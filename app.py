# ============================================================================
# HARMONY API + DASHBOARD  (FastAPI)
# ----------------------------------------------------------------------------
# CHALANE KA TARIQA:
#   uvicorn app:app --reload
# Phir browser mein kholein:  http://127.0.0.1:8000
#
# YE terminal input() ki jagah hai. Client web par kaam karega:
#   1. History add karo (text ya server-path)
#   2. Related past data search karo
#   3. Draft submit karo -> report + rating
#   4. Approve/Reject buttons (Approve par safe version banti hai + history mein save)
# ============================================================================

import json
import uuid
from fastapi import FastAPI, Depends, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel

import harmony   # saara agent + Qdrant logic yahan se aata hai

# --- Enterprise backend: DB + Auth + Org + Docs + Security + Billing + SSO ---
from datetime import datetime
from sqlmodel import Session
from sqlalchemy import text as _sa_text
from db import init_db, engine, User, Document, Review, HistoryItem
import routes_auth, routes_org, routes_documents, routes_security, routes_billing, routes_sso, routes_admin
from routes_auth import audit
from auth import current_user, current_org_id, require_role, user_membership

app = FastAPI(title="Harmony — Disclosure Consistency Copilot")

# Next.js frontend runs on a different origin/port in dev (and possibly prod) —
# allow it to call this API directly with a Bearer token.
from fastapi.middleware.cors import CORSMiddleware
import os as _os_cors
_cors_origins = [o.strip() for o in _os_cors.getenv(
    "CORS_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000,http://localhost:3100,http://127.0.0.1:3100"
).split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

init_db()

# --- Email configuration, checked once at startup ---
# Verification is skipped when SMTP is unconfigured, so a wrong setting means
# signup keeps working while no code is ever delivered. Say so loudly here
# rather than letting it be discovered by a customer who never got their code.
import emailer as _emailer
_EMAIL_PROBLEM = _emailer.config_problem()
if _EMAIL_PROBLEM:
    print(f"EMAIL: {_EMAIL_PROBLEM}")
    _emailer.verify_connection()      # sets the healthy flag to False
else:
    _ok, _detail = _emailer.verify_connection()
    _EMAIL_PROBLEM = None if _ok else _detail
    if not _ok:
        print(f"EMAIL: {_emailer.transport()} check FAILED - {_detail}")
if _EMAIL_PROBLEM:
    print("EMAIL: verification codes will NOT be sent; signup proceeds WITHOUT verifying "
          "so users are not locked out of a product whose code can never arrive.")
else:
    print(f"EMAIL: {_emailer.transport()} ok - verification codes will be sent.")
for r in (routes_auth.router, routes_org.router, routes_documents.router,
          routes_security.router, routes_billing.router, routes_sso.router,
          routes_admin.router):
    app.include_router(r)

# --- Unhandled errors ---
# Dev: full traceback in the response, because that is how you debug quickly.
# Production: the traceback goes to the LOGS ONLY. Returning it to the client
# hands anyone who can trigger a 500 your file paths, library versions and
# often the surrounding code — a free map of the system. The client gets an
# error id instead, which you can grep for in the Render logs.
import traceback as _tb
import uuid as _uuid_err
from fastapi import Request as _Req
from fastapi.responses import JSONResponse as _JR
from auth import APP_ENV as _APP_ENV


@app.exception_handler(Exception)
async def _show_error(request: _Req, exc: Exception):
    error_id = _uuid_err.uuid4().hex[:12]
    tb = "".join(_tb.format_exception(type(exc), exc, exc.__traceback__))
    print(f"[error {error_id}] {request.method} {request.url.path}")
    print(tb)

    if _APP_ENV == "production":
        return _JR(status_code=500, content={
            "detail": "Something went wrong on our side. Quote this reference if you contact support.",
            "error_id": error_id,
        })
    return _JR(status_code=500, content={"detail": tb[-1400:], "error_id": error_id})

# pending reviews (transient — jaan-boojh kar permanent nahi rakhte, taake koi
# "known contradiction" ka permanent record na bane). review_id -> result
PENDING: dict = {}


# ---- request models ----
class IngestText(BaseModel):
    text: str
    company: str | None = "Unknown"


class IngestPath(BaseModel):
    path: str


class ReviewIn(BaseModel):
    draft: str


class DecisionIn(BaseModel):
    review_id: str
    decision: str   # "approve" ya "reject"


# ---- API endpoints ----

def _record_history_item(org_id: str, user_id: str, company: str, source_file: str, doc_type: str, chunk_count: int):
    if chunk_count <= 0:
        return
    with Session(engine) as s:
        s.add(HistoryItem(org_id=org_id, company=company or "Unknown", source_file=source_file,
                          doc_type=doc_type, chunk_count=chunk_count, added_by=user_id))
        s.commit()


@app.post("/api/ingest_text")
def ingest_text(body: IngestText, user: User = Depends(require_role("editor")),
                org_id: str = Depends(current_org_id)):
    harmony.set_org(org_id)   # org-scoped tag
    n = harmony.add_to_history(body.text, company=body.company or "Unknown", doc_type="history")
    _record_history_item(org_id, user.id, body.company or "Unknown", "", "history", n)
    audit(org_id, user.id, "history.ingested", "text")
    return {"status": "ok", "detail": "Statement added to history."}


@app.post("/api/ingest_path")
def ingest_path(body: IngestPath, user: User = Depends(require_role("editor")),
                org_id: str = Depends(current_org_id)):
    harmony.set_org(org_id)
    res = harmony.ingestion_agent({"messages": [{"role": "user", "content": body.path}]})
    audit(org_id, user.id, "history.ingested", body.path)
    return {"status": "ok", "detail": res["messages"][-1]["content"]}


@app.get("/api/search")
def search(q: str, user: User = Depends(current_user), org_id: str = Depends(current_org_id)):
    harmony.set_org(org_id)   # sirf is org ki history search ho
    return {"related": harmony.retrieve_context(q, k=5)}


def _extract_text(filename: str, data: bytes) -> str:
    # ek file (pdf/txt/docx) se plain text nikaalta hai
    import io
    name = (filename or "").lower()
    if name.endswith(".txt"):
        return data.decode("utf-8-sig", errors="replace")   # utf-8-sig strips a leading BOM if present
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if name.endswith(".docx"):
        from docx import Document as _Docx
        d = _Docx(io.BytesIO(data))
        return "\n".join(p.text for p in d.paragraphs)
    return ""


@app.post("/api/upload")
async def upload(files: list[UploadFile] = File(...),
                 user: User = Depends(require_role("editor")),
                 org_id: str = Depends(current_org_id)):
    # ek ya kai files (folder) -> text nikaal kar wapas do (frontend draft mein daal dega)
    parts = []
    skipped = []
    for f in files:
        data = await f.read()
        text = _extract_text(f.filename, data)
        if text.strip():
            parts.append(f"### {f.filename}\n{text}")
        else:
            skipped.append(f.filename)
    combined = "\n\n".join(parts)
    audit(org_id, user.id, "file.uploaded", ", ".join(f.filename for f in files)[:200])
    return {"count": len(parts), "skipped": skipped,
            "text": combined, "words": len(combined.split())}


@app.post("/api/upload_history")
async def upload_history(files: list[UploadFile] = File(...), company: str = Form("Unknown"),
                         user: User = Depends(require_role("editor")),
                         org_id: str = Depends(current_org_id)):
    # ek ya kai files (folder) -> har file do jagah jata hai:
    #   1) evidence history (Qdrant) -> future reviews isse compare karenge
    #   2) Document record (DB) -> "My Documents" mein bhi ek "previous draft" ke tor par dikhta hai
    # (koi AI review nahi chalta — ye already-published/historical files hain, review ki zaroorat nahi.)
    harmony.set_org(org_id)
    added = []
    skipped = []
    for f in files:
        data = await f.read()
        text = _extract_text(f.filename, data)
        if text.strip():
            n = harmony.add_to_history(text, company=company or "Unknown", doc_type="history", source_file=f.filename)
            _record_history_item(org_id, user.id, company or "Unknown", f.filename, "history", n)
            title = (f.filename.rsplit(".", 1)[0] if f.filename else "Untitled")
            with Session(engine) as s:
                s.add(Document(org_id=org_id, title=title, doc_type="Historical Disclosure",
                               content=text, status="Published", risk="Low", created_by=user.id))
                s.commit()
            added.append(f.filename)
        else:
            skipped.append(f.filename)
    audit(org_id, user.id, "history.ingested", ", ".join(added)[:200])
    return {"count": len(added), "added": added, "skipped": skipped}


def _risk_from_issues(issues: list) -> str:
    severities = {i.get("severity") for i in issues}
    if "high" in severities:
        return "High"
    if "medium" in severities:
        return "Medium"
    return "Low"


def _persist_review(result: dict, org_id: str, user_id: str) -> dict:
    # run_review() / run_review_stream() dono ka result yahan persist hota hai (shared,
    # taake dono endpoints ek hi logic use karein — do jagah duplicate na ho).
    review_id = str(uuid.uuid4())
    issues = result.get("issues", [])
    evidence = result.get("evidence", [])
    risk = _risk_from_issues(issues)
    doc_type = (result.get("draft_topic") or "").strip().title() or "Disclosure Draft"
    title = (result["company"] or "Draft") + " — " + doc_type
    with Session(engine) as s:
        doc = Document(org_id=org_id, title=title, doc_type=doc_type,
                       content=result["draft_text"], status="In Review", risk=risk, created_by=user_id)
        s.add(doc); s.commit(); s.refresh(doc)
        doc_id = doc.id   # session ke andar hi capture karo (commit ke baad attributes expire ho jate hain)
        rev = Review(id=review_id, org_id=org_id, document_id=doc_id, company=result["company"],
                     average_rating=result["average_rating"], critic_verdict=result["critic_verdict"],
                     report=result["final_summary"], issues_json=json.dumps(issues),
                     evidence_json=json.dumps(evidence),
                     status="pending", created_by=user_id)
        s.add(rev); s.commit()
    PENDING[review_id] = {**result, "document_id": doc_id}
    return {
        "review_id": review_id,
        "document_id": doc_id,
        "company": result["company"],
        "doc_type": doc_type,
        "risk": risk,
        "average_rating": result["average_rating"],
        "critic_verdict": result["critic_verdict"],
        "report": result["final_summary"],
        "issues": issues,
        "evidence": evidence,
    }


@app.post("/api/review")
def review(body: ReviewIn, user: User = Depends(require_role("editor")),
           org_id: str = Depends(current_org_id)):
    # login + editor+ role zaroori. Org-scoped review (sirf is org ki history se compare).
    result = harmony.run_review(body.draft, org_id=org_id)
    payload = _persist_review(result, org_id, user.id)
    audit(org_id, user.id, "review.created", result["company"])
    return payload


@app.post("/api/review_stream")
def review_stream(body: ReviewIn, user: User = Depends(require_role("editor")),
                  org_id: str = Depends(current_org_id)):
    # Wahi review pipeline, lekin Server-Sent Events ke zariye live progress + report ke
    # tokens stream karta hai. Aakhri "done" event mein wahi payload hota hai jo /api/review
    # deta hai, taake frontend seedha Review Workspace render kar sake.
    user_id = user.id

    def gen():
        try:
            for event in harmony.run_review_stream(body.draft, org_id=org_id):
                if event.get("type") == "done":
                    payload = _persist_review(event["result"], org_id, user_id)
                    audit(org_id, user_id, "review.created", event["result"].get("company", ""))
                    yield f"data: {json.dumps({'type': 'done', 'result': payload})}\n\n"
                else:
                    yield f"data: {json.dumps(event)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.post("/api/decision")
def decision(body: DecisionIn, user: User = Depends(require_role("reviewer")),
             org_id: str = Depends(current_org_id)):
    """Approve or reject a review. REVIEWER+ only, org-scoped.

    Reads the review from the DATABASE, not from the in-memory PENDING dict.

    PENDING did not survive a restart, and a free-tier service sleeps after
    fifteen minutes - so every approve and reject returned 404 "Review not found
    or already resolved" while the review sat perfectly intact in Postgres. Any
    redeploy did the same. It only ever worked in a process that had served the
    review itself.

    It was also a tenant leak: PENDING was keyed by review id alone, with no org
    check, so a reviewer in one workspace could decide another workspace's
    review by id. Loading through the org filter closes that.
    """
    with Session(engine) as s:
        rev = s.get(Review, body.review_id)
        # The org filter is the access check, not a nicety.
        if not rev or rev.org_id != org_id:
            return JSONResponse({"error": "Review not found."}, status_code=404)
        if rev.status != "pending":
            return JSONResponse(
                {"error": f"This review was already {rev.status}."}, status_code=409)
        try:
            issues = json.loads(rev.issues_json or "[]")
        except (json.JSONDecodeError, TypeError):
            issues = []
        try:
            evidence = json.loads(rev.evidence_json or "[]")
        except (json.JSONDecodeError, TypeError):
            evidence = []
        result = {
            "company": rev.company,
            "final_summary": rev.report,
            "average_rating": rev.average_rating,
            "critic_verdict": rev.critic_verdict,
            "issues": issues,
            "evidence": evidence,
            "document_id": rev.document_id,
        }

    def _update_status(new_status):
        with Session(engine) as s:
            rev = s.get(Review, body.review_id)
            if rev:
                rev.status = new_status; rev.decided_by = user.id; rev.decided_at = datetime.utcnow()
                s.add(rev)
                doc = s.get(Document, rev.document_id) if rev.document_id else None
                if doc:
                    doc.status = "Published" if new_status == "approved" else "Changes Requested"
                    s.add(doc)
                s.commit()

    if body.decision == "approve":
        publish_result = harmony.publish_review(result, org_id=org_id)
        _update_status("approved")
        PENDING.pop(body.review_id, None)
        audit(org_id, user.id, "review.approved", result.get("company", ""))
        _record_history_item(org_id, user.id, result.get("company", "Unknown"), "", "approved",
                             publish_result.get("chunks_added", 0))
        return {"status": "approved", "final_version": publish_result["final_version"]}

    _update_status("rejected")
    PENDING.pop(body.review_id, None)
    audit(org_id, user.id, "review.rejected", result.get("company", ""))
    return {"status": "rejected", "detail": "Review discarded. Nothing was saved."}


@app.get("/healthz", tags=["ops"])
def healthz():
    # Render/Railway is path ko baar baar ping karte hain. Sasta rakho: "/" 85KB HTML
    # return karta hai, wo health check ke liye fizool hai. DB ko ek halka touch dete
    # hain taake "up but database gone" wali halat bhi pakri jaye.
    db_ok = True
    try:
        with Session(engine) as s:
            s.exec(_sa_text("SELECT 1"))
    except Exception:
        db_ok = False

    # Report the DIALECT, not just reachability. SELECT 1 succeeds on SQLite too,
    # so a plain "up" hid the worst possible misconfiguration: no APP_DATABASE_URL,
    # a SQLite file on an ephemeral container disk, and every account silently
    # lost on the next restart or deploy.
    dialect = engine.dialect.name
    ephemeral = dialect == "sqlite" and _APP_ENV == "production"

    body = {
        "status": "ok" if db_ok and not ephemeral and not _EMAIL_PROBLEM else "degraded",
        "database": "up" if db_ok else "down",
        "engine": dialect,
        "email": "ok" if not _EMAIL_PROBLEM else "unavailable",
        "email_transport": _emailer.transport(),
    }
    if _EMAIL_PROBLEM:
        body["email_detail"] = _EMAIL_PROBLEM
    if ephemeral:
        body["warning"] = ("Running on SQLite in production. This filesystem is ephemeral - "
                           "every account and review is lost on the next restart or deploy. "
                           "Set APP_DATABASE_URL to a Postgres connection string.")
    return JSONResponse(body, status_code=200 if db_ok else 503)


import os as _os

# ----------------------------------------------------------------------------
# THE UI
# ----------------------------------------------------------------------------
# harmony-web is built as a static export (next.config.mjs: output "export") and
# served from this same service. One deployment instead of two, and because the
# UI and the API share an origin there is no CORS to configure or get wrong -
# which was the single most likely thing to break in a split deployment.
#
# There is no second UI any more. The old single-file dashboard used to answer
# here as a fallback, and having two live interfaces was actively harmful: they
# drifted, and a feature added to one silently did not exist in the other. If
# the export is missing, that means the Docker Node stage did not run, and the
# right response is to say so plainly rather than quietly serve something else.
_HERE = _os.path.dirname(__file__)
_WEB_DIR = _os.path.join(_HERE, "harmony-web", "out")
_HAS_WEB = _os.path.isdir(_WEB_DIR) and _os.path.isfile(_os.path.join(_WEB_DIR, "index.html"))

_NO_UI_PAGE = """<!doctype html>
<meta charset="utf-8"><title>Harmony - UI not built</title>
<style>body{background:#08090a;color:#f2f2f3;font:14px/1.6 system-ui,sans-serif;
display:grid;place-items:center;min-height:100vh;margin:0}
div{max-width:460px;padding:28px}code{background:#16181b;padding:2px 6px;border-radius:5px}
a{color:#8f9aa8}</style>
<div>
  <h1 style="font-size:19px;font-weight:600;margin:0 0 10px">The interface was not built</h1>
  <p style="color:#8b8f96">The API is running, but <code>harmony-web/out</code> is missing, so
  there is nothing to serve at this address.</p>
  <p style="color:#8b8f96">Locally, run <code>npm run build</code> in <code>harmony-web/</code>.
  On a deploy, the Docker Node stage did not complete - check the build log.</p>
  <p style="color:#8b8f96">The API itself is unaffected: <a href="/docs">/docs</a> ·
  <a href="/healthz">/healthz</a></p>
</div>"""


if not _HAS_WEB:
    print("harmony-web/out not found - the product UI will NOT be served. "
          "Run `npm run build` in harmony-web/, or check the Docker Node stage.")

    @app.get("/", response_class=HTMLResponse)
    def home():
        return HTMLResponse(_NO_UI_PAGE, status_code=503)


# ----------------------------------------------------------------------------
# Static UI mount - MUST be last.
# ----------------------------------------------------------------------------
# A mount at "/" shadows every route declared after it, so this sits at the end
# of the file, below every API route. html=True makes StaticFiles resolve a
# directory to its index.html, which is how the export's deep links
# (/app/documents/ -> app/documents/index.html) work without any rewrite rules.
if _HAS_WEB:
    from fastapi.staticfiles import StaticFiles

    app.mount("/", StaticFiles(directory=_WEB_DIR, html=True), name="harmony-web")
